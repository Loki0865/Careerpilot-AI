import os
import sys
import docx

# Add path for imports
sys.path.append(os.path.abspath(os.path.dirname(__file__)))

from security.safe_file_access import SAFE_UPLOAD_DIR

def create_sample_resume():
    doc = docx.Document()
    doc.add_heading('JANE DOE', 0)
    
    doc.add_heading('Professional Summary', level=1)
    p1 = doc.add_paragraph(
        "Motivated Junior Software Engineer looking for a backend role. "
        "Experienced in Python and basic web technologies. Good team player."
    )
    
    doc.add_heading('Experience', level=1)
    
    # TechCorp Experience
    doc.add_heading('TechCorp - Junior Software Engineer (2025 - Present)', level=2)
    doc.add_paragraph("Wrote python code for backend projects. Fixed bugs reported by QA team. Worked with team members. Helped users with login issues.")
    
    # WebInc Experience
    doc.add_heading('WebInc - Software Engineering Intern (2024 - 2025)', level=2)
    doc.add_paragraph("Created web pages using HTML and CSS. Used git for version control. Fixed alignment of elements on the main dashboard.")
    
    doc.add_heading('Skills', level=1)
    doc.add_paragraph("Python, HTML, CSS, JavaScript, Git, Windows, Microsoft Office")
    
    doc.add_heading('Education', level=1)
    doc.add_paragraph("Bachelor of Science in Computer Science, State University, 2024")
    
    filepath = os.path.join(SAFE_UPLOAD_DIR, "jane_doe_resume.docx")
    doc.save(filepath)
    print(f"Created sample resume at: {filepath}")

def create_sample_jd():
    doc = docx.Document()
    doc.add_heading('Python Backend Engineer - Job Description', 0)
    
    doc.add_heading('Company Overview', level=1)
    doc.add_paragraph("We are a fast-growing tech company building AI-powered financial solutions.")
    
    doc.add_heading('Role Overview', level=1)
    doc.add_paragraph(
        "We are looking for a Python Backend Engineer with 2+ years of experience to join our core API team. "
        "You will build secure, scalable backend services, optimize data queries, and deploy cloud resources."
    )
    
    doc.add_heading('Required Technical Skills', level=1)
    doc.add_paragraph("- Strong proficiency in Python and FastAPI or Flask frameworks.")
    doc.add_paragraph("- Experience with relational databases like PostgreSQL or MySQL.")
    doc.add_paragraph("- Experience with Docker containerization.")
    doc.add_paragraph("- Basic familiarity with AWS services (EC2, S3, RDS).")
    doc.add_paragraph("- Proficient with Git and RESTful API design.")
    
    doc.add_heading('Responsibilities', level=1)
    doc.add_paragraph("- Design and implement scalable REST APIs.")
    doc.add_paragraph("- Optimize database queries to ensure high performance.")
    doc.add_paragraph("- Write clean, testable, and well-documented Python code.")
    
    filepath = os.path.join(SAFE_UPLOAD_DIR, "backend_job_description.docx")
    doc.save(filepath)
    print(f"Created sample JD at: {filepath}")

if __name__ == "__main__":
    create_sample_resume()
    create_sample_jd()
